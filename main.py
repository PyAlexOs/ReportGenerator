import docx
from docx.enum.text import (WD_PARAGRAPH_ALIGNMENT,
                            WD_BREAK_TYPE,
                            WD_COLOR_INDEX,
                            WD_LINE_SPACING,
                            WD_TAB_ALIGNMENT,
                            WD_TAB_LEADER,
                            WD_UNDERLINE)
from docx.enum.style import (WD_BUILTIN_STYLE,
                             WD_STYLE_TYPE)
from docx.enum.dml import (MSO_COLOR_TYPE,
                           MSO_THEME_COLOR_INDEX)
from docx.shared import Pt, Cm, Mm, RGBColor
from enum import Enum
from os import path
from sys import argv
import json
import re


class Report:
    """ Provides methods for working with the generated report in accordance with the state standard """

    def __init__(self, filename: str, styles_path="styles.json"):
        self.styles_path = path.abspath(styles_path)
        self.filename = filename
        self.check_filename()

        try:
            document = docx.Document()
            document.save(filename)
            self.document = docx.Document(filename)
        except PermissionError:
            exit("To generate a report, close the file that you specified for recording")

    def __call__(self):
        self.set_styles()
        self.parse_content()
        self.document.save(self.filename)

    def check_filename(self):
        """ Checks if the file path and file name are correct """
        if re.match(r"^(/+|\\+)$", self.filename) is not None:
            exit("Wrong file path")

        dlm: str = '/' if self.filename.find("/") != -1 else "\\"
        if not path.exists(dlm.join(self.filename.split(dlm)[:-1:])):
            exit("File path is specified incorrectly")

        if len(self.filename.split(".")) < 2 or self.filename.split(".")[-1] != "docx":
            exit("Incorrect file format")

        """if path.exists(filename):
            if input("File was found at the specified path. Do you really want to overwrite it? (Y/N)\n") not in ["Y", "y"]:
                exit()"""

    def set_styles(self):
        """ Adds all the necessary styles from .json file"""
        if not path.exists(self.styles_path):
            exit("The file containing the styles was not found")

        with open(self.styles_path, "r", encoding="utf-8") as file:
            styles_from_file = json.load(file)

        styles = self.document.styles
        for style in styles:
            styles[style.name].delete()

        for (i, (name, values)) in enumerate(styles_from_file["styles"].items()):
            try:
                current_style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
                font = current_style.font
                font.name = values["font"]["name"]
                font.size = Pt(values["font"]["size"])

                font.all_caps = values["font"]["all_caps"]
                font.bold = values["font"]["bold"]
                font.italic = values["font"]["italic"]
                font.underline = values["font"]["underline"]

                font.color.ColorFormat = MSO_COLOR_TYPE.RGB
                font.color.rgb = RGBColor(*values["font"]["color"])

                font.math = values["font"]["math"]
                font.no_proof = values["font"]["no_proof"]

                paragraph = current_style.paragraph_format
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT(values["paragraph_format"]["alignment"])
                paragraph.first_line_indent = Cm(values["paragraph_format"]["first_line_indent"])
                paragraph.line_spacing_rule = WD_LINE_SPACING(values["paragraph_format"]["line_spacing_rule"])

                paragraph.left_indent = Cm(values["paragraph_format"]["left_indent"])
                paragraph.right_indent = Cm(values["paragraph_format"]["right_indent"])
                paragraph.space_before = Mm(values["paragraph_format"]["space_before"])
                paragraph.space_after = Mm(values["paragraph_format"]["space_after"])

                paragraph.keep_together = values["paragraph_format"]["keep_together"]
                paragraph.keep_with_next = values["paragraph_format"]["keep_with_next"]
                paragraph.page_break_before = values["paragraph_format"]["page_break_before"]
                paragraph.widow_control = values["paragraph_format"]["widow_control"]
                paragraph.next_paragraph_style = values["paragraph_format"]["next_paragraph_style"]

            except (TypeError, KeyError) as error:
                exit(f"Incorrect style attribute: {error}")

    def parse_content(self):
        """ Parses files in the directory and composes the document from them """
        self.document.add_paragraph("Heading1 but normal", "Normal")
        """self.document.add_paragraph("Heading 2", style="Heading 2")
        self.document.add_paragraph("Heading 3", style="Heading 3")
        self.document.add_paragraph("NORMal Text herrrre", style="Normal")"""


class ParagraphType(Enum):
    HEADING_1 = 1
    HEADING_2 = 2
    HEADING_3 = 3
    BODY = 4

    PICTURE = 10
    TABLE = 11
    FORMULA = 12
    LISTING = 13


# python main.py documents/testReport.docx
def main():
    try:
        filename = argv[1]
    except IndexError:
        filename = input('Enter file name: ')

    report = Report(path.abspath(filename))
    report()

    token_list: list[list[str, ParagraphType]] = list()
    with open("documents/test.md", "r", encoding="utf-8") as file:
        lines = file.readlines()
        for (i, line) in enumerate(lines):
            if line.startswith("###"):
                token_list.append(line)


if __name__ == '__main__':
    main()
